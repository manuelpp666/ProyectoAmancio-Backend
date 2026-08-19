import os
import re
import asyncio
import threading
import queue as thread_queue
import time
import requests
from dotenv import load_dotenv
from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException, status
from google import genai
from google.genai import types as genai_types
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document as LangDocument
from typing import List, Tuple, AsyncGenerator
from datetime import datetime
import uuid
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from app.db.database import get_db
from .models import Chatbot
from .schemas import ChatbotResponse
from app.core.util.security import get_current_user
from tenacity import retry, stop_after_attempt, wait_random_exponential, retry_if_exception_type

# --- LECTOR DE PDF CON SOPORTE DE TABLAS ---
# pip install pdfplumber
import pdfplumber

# --- OCR (solo se activa si una página no tiene texto extraíble, es decir, si
# es una imagen escaneada). Esto evita gastar OCR en documentos normales y
# mantiene el costo bajo. ---
#
# No basta con que los paquetes de Python se importen: pytesseract y pdf2image
# son envoltorios que llaman a dos programas externos (Tesseract y Poppler). Si
# se importan pero los programas no están, el OCR fallaba al procesar el primer
# PDF escaneado en vez de saltárselo. Por eso aquí se comprueba que los
# ejecutables existan de verdad y que esté el idioma español, y solo entonces se
# enciende el OCR.
#
# Las rutas se pueden fijar en el .env (TESSERACT_CMD, POPPLER_PATH,
# TESSDATA_PREFIX). Si no, se buscan en el PATH y en las rutas donde los deja
# una instalación normal en Windows, para no depender de que el PATH esté bien
# en la consola desde la que se arranque el servidor.
import shutil

def _buscar_ejecutable(nombre_env, ejecutable, candidatos):
    ruta = os.getenv(nombre_env)
    if ruta and os.path.exists(ruta):
        return ruta
    encontrado = shutil.which(ejecutable)
    if encontrado:
        return encontrado
    for c in candidatos:
        c = os.path.expandvars(c)
        if os.path.exists(c):
            return c
    return None


TESSERACT_CMD = _buscar_ejecutable(
    "TESSERACT_CMD", "tesseract",
    [r"%ProgramFiles%\Tesseract-OCR\tesseract.exe",
     r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"],
)
_pdftoppm = _buscar_ejecutable(
    "POPPLER_PDFTOPPM", "pdftoppm",
    [r"%LOCALAPPDATA%\Microsoft\WinGet\Packages"
     r"\oschwartz10612.Poppler_Microsoft.Winget.Source_8wekyb3d8bbwe"
     r"\poppler-25.07.0\Library\bin\pdftoppm.exe"],
)
POPPLER_PATH = os.getenv("POPPLER_PATH") or (os.path.dirname(_pdftoppm) if _pdftoppm else None)

# Carpeta de idiomas. La del sistema (Program Files) solo trae inglés y añadir
# el español ahí exige permisos de administrador, así que el proyecto lleva la
# suya en ProyectoAmancio-Backend/tessdata.
#
# Tiene que estar dentro del proyecto, NO en AppData: si Python es el de
# Microsoft Store, Windows virtualiza sus escrituras en AppData y la carpeta
# queda en un almacén privado del paquete que Tesseract (aplicación normal) no
# puede ver. Con la carpeta junto al proyecto, ambos ven lo mismo.
_TESSDATA_PROYECTO = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))),
    "tessdata",
)
_tessdata = os.getenv("TESSDATA_PREFIX") or _TESSDATA_PROYECTO
if os.path.isdir(_tessdata):
    os.environ["TESSDATA_PREFIX"] = _tessdata

# El OCR se prepara la PRIMERA VEZ que hace falta, no al arrancar.
#
# Antes esto se importaba aquí mismo, y con ello entraba numpy en el arranque
# del servidor. En el hosting compartido eso tumbaba la aplicación entera: numpy
# carga OpenBLAS, OpenBLAS pide un hilo por núcleo (32 en ese servidor), el plan
# no permite tantos y el proceso moría antes de levantar la API. El campus se
# quedaba sin backend por una función que casi nunca se usa.
#
# Ahora el arranque no toca numpy. Solo se paga ese coste si de verdad llega un
# PDF escaneado, y el resultado se recuerda para no repetir la comprobación.
_ocr_estado = None          # None = sin comprobar todavía
_pytesseract = None
_convert_from_path = None


def _preparar_ocr() -> bool:
    """Comprueba e inicializa el OCR una sola vez. True si se puede usar."""
    global _ocr_estado, _pytesseract, _convert_from_path
    if _ocr_estado is not None:
        return _ocr_estado

    _ocr_estado = False
    try:
        import pytesseract
        from pdf2image import convert_from_path

        if TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

        faltan = []
        if not TESSERACT_CMD:
            faltan.append("Tesseract (programa)")
        if not POPPLER_PATH:
            faltan.append("Poppler (programa)")
        if not faltan:
            idiomas = set(pytesseract.get_languages(config=""))
            if "spa" not in idiomas:
                faltan.append("idioma español de Tesseract (spa.traineddata)")

        if faltan:
            print("[WARN] OCR desactivado, falta: " + ", ".join(faltan))
        else:
            _pytesseract = pytesseract
            _convert_from_path = convert_from_path
            _ocr_estado = True
            print(f"[OK] OCR activo · Tesseract: {TESSERACT_CMD} · idiomas: spa+eng")
    except ImportError:
        print(
            "[WARN] pytesseract/pdf2image no están instalados. El fallback de OCR para "
            "PDFs escaneados estará desactivado. Instala con:\n"
            "  pip install pytesseract pdf2image"
        )
    except Exception as e:  # p. ej. Tesseract instalado pero roto
        print(f"[WARN] OCR desactivado por un error al comprobarlo: {e}")

    return _ocr_estado


load_dotenv()

FILE_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(FILE_DIR)))
UPLOAD_DIR = os.path.join(BASE_DIR, "media", "chatbot_files")
os.makedirs(UPLOAD_DIR, exist_ok=True)

MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_DOCUMENTS = 5

NAMESPACE = "colegio"

# Debe coincidir EXACTO con la dimensión configurada en tu índice de Pinecone.
# gemini-embedding-2 soporta dimensión variable (output_dimensionality); si
# algún día recreas el índice con otra dimensión, este es el único lugar que
# hay que tocar -- se usa tanto para indexar (embed_documents) como para
# consultar (embed_query y get_native_embedding_3072), así que ambos lados
# quedan sincronizados automáticamente.
EMBEDDING_DIMENSIONS = 3072

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
client = genai.Client(api_key=GOOGLE_API_KEY)
INDEX_NAME = os.getenv("INDEX_NAME", "colegio-knowledge")

os.environ["PINECONE_API_KEY"] = PINECONE_API_KEY

router = APIRouter(prefix="/chatbot", tags=["Chatbot"])

class NativeGeminiEmbeddings:
    """Genera embeddings usando el cliente REST de google-genai en vez de
    gRPC (ver por qué en el comentario de más abajo). Ya no hereda de
    langchain_core.embeddings.Embeddings porque no se usa con
    PineconeVectorStore -- ese SDK también quedó fuera, ver PineconeREST."""

    def __init__(self, genai_client: "genai.Client", model: str = "models/gemini-embedding-2"):
        self._client = genai_client
        self._model = model

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        # Secuencial a propósito: nada de ThreadPoolExecutor/hilos aquí.
        # Es un poco más lento que en paralelo, pero es justamente la
        # concurrencia con hilos la que rompe en cPanel.
        vectors: List[List[float]] = []
        for text in texts:
            response = self._client.models.embed_content(
                model=self._model,
                contents=text,
                config=genai_types.EmbedContentConfig(
                    task_type="RETRIEVAL_DOCUMENT",
                    output_dimensionality=EMBEDDING_DIMENSIONS,
                ),
            )
            vectors.append(response.embeddings[0].values)
        return vectors

    def embed_query(self, text: str) -> List[float]:
        response = self._client.models.embed_content(
            model=self._model,
            contents=text,
            config=genai_types.EmbedContentConfig(
                task_type="RETRIEVAL_QUERY",
                output_dimensionality=EMBEDDING_DIMENSIONS,
            ),
        )
        return response.embeddings[0].values


embeddings = NativeGeminiEmbeddings(client)


# ==========================================================================
# CLIENTE PINECONE 100% REST (reemplaza al SDK oficial `pinecone-client` /
# `langchain_pinecone`) -- SOLUCIÓN DEFINITIVA AL "can't start new thread" /
# "'DummyProcess' object has no attribute 'terminate'"
# ==========================================================================
# El SDK oficial de Pinecone arma, para manejar sus conexiones HTTP
# concurrentes, un multiprocessing.pool.ThreadPool (esa clase usa por dentro
# multiprocessing.dummy.Process, que en el error aparece como "DummyProcess").
# Eso es lo tercero (después de OpenBLAS y gRPC) que intenta abrir hilos del
# sistema operativo en un entorno cPanel/CloudLinux donde la cuenta ya tiene
# su cupo de hilos (LVE) casi agotado -- por eso truena distinto según cuánto
# margen quedaba en ese momento (a veces "DummyProcess...terminate", a veces
# derecho "can't start new thread").
#
# La única forma de eliminar el riesgo por completo es dejar de usar el SDK
# para las llamadas de datos (upsert/fetch/delete/query) y hablarle a la API
# REST de Pinecone directamente con `requests`, que NO abre hilos propios:
# reutiliza sockets dentro del mismo hilo que hace la llamada (igual que ya
# hace el cliente REST de Gemini, que nunca ha fallado en cPanel).
class PineconeREST:
    API_VERSION = "2025-10"

    def __init__(self, api_key: str, index_name: str):
        self._session = requests.Session()
        self._session.headers.update({
            "Api-Key": api_key,
            "Content-Type": "application/json",
            "X-Pinecone-Api-Version": self.API_VERSION,
        })
        self._host = self._resolve_host(index_name)

    def _resolve_host(self, index_name: str) -> str:
        # Se resuelve UNA sola vez al arrancar el servidor (plano de control),
        # no en cada request -- esto nunca es el cuello de botella.
        resp = self._session.get(f"https://api.pinecone.io/indexes/{index_name}", timeout=15)
        resp.raise_for_status()
        host = resp.json().get("host")
        if not host:
            raise RuntimeError(f"No se pudo resolver el host del índice '{index_name}' en Pinecone.")
        return f"https://{host}"

    def upsert(self, vectors: List[dict], namespace: str) -> None:
        resp = self._session.post(
            f"{self._host}/vectors/upsert",
            json={"vectors": vectors, "namespace": namespace},
            timeout=30,
        )
        resp.raise_for_status()

    def fetch_existing_ids(self, ids: List[str], namespace: str) -> set:
        """Si la verificación misma falla (red, timeout, etc.), se asume que
        faltan todos los IDs -- eso fuerza un reintento en vez de reportar
        un falso "ya está" que ocultaría chunks realmente perdidos."""
        if not ids:
            return set()
        try:
            resp = self._session.get(
                f"{self._host}/vectors/fetch",
                params={"ids": ids, "namespace": namespace},
                timeout=15,
            )
            resp.raise_for_status()
            return set(resp.json().get("vectors", {}).keys())
        except Exception as e:
            print(f"[PINECONE][WARN] No se pudo verificar contra Pinecone: {e}")
            return set()

    def delete(self, ids: List[str], namespace: str) -> None:
        if not ids:
            return
        resp = self._session.post(
            f"{self._host}/vectors/delete",
            json={"ids": ids, "namespace": namespace},
            timeout=15,
        )
        resp.raise_for_status()

    def query(self, vector: List[float], top_k: int, namespace: str) -> List[dict]:
        resp = self._session.post(
            f"{self._host}/query",
            json={
                "vector": vector,
                "topK": top_k,
                "namespace": namespace,
                "includeMetadata": True,
            },
            timeout=20,
        )
        resp.raise_for_status()
        return resp.json().get("matches", [])


pinecone_client = PineconeREST(PINECONE_API_KEY, INDEX_NAME)

# ==========================================================================
# 1. PROTECCIÓN DE DATOS SENSIBLES
# ==========================================================================
# Filtro conservador: solo redacta patrones "etiquetados" (ej. "DNI: 12345678")
# para no destruir tablas de pagos legítimas que también tienen números.
# Los patrones de tarjeta/email SÍ se redactan siempre porque su formato es
# lo bastante distintivo como para no generar falsos positivos.

SENSITIVE_PATTERNS = {
    "DNI": re.compile(r"(?i)\bDNI\s*[:\-]?\s*(\d{8})\b"),
    "RUC": re.compile(r"(?i)\bRUC\s*[:\-]?\s*(\d{11})\b"),
    "TELEFONO": re.compile(r"(?i)\b(?:tel[eé]fono|cel(?:ular)?|whats\s*app)\s*[:\-]?\s*(\+?51[\s\-]?)?(9\d{2}[\s\-]?\d{3}[\s\-]?\d{3})\b"),
    "CUENTA_BANCARIA": re.compile(r"(?i)\b(?:cuenta|cci)\s*(?:bancaria|interbancaria)?\s*[:\-]?\s*(\d[\d\s\-]{9,20}\d)\b"),
    "TARJETA": re.compile(r"\b(?:\d{4}[\s\-]){3}\d{4}\b"),
    "EMAIL": re.compile(r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b"),
    "CONTRASENA": re.compile(r"(?i)\b(?:contrase[nñ]a|password|clave)\s*[:\-]?\s*(\S{4,})\b"),
}


def redact_sensitive_data(text: str) -> Tuple[str, dict]:
    """
    Reemplaza datos sensibles etiquetados por un placeholder ANTES de que
    lleguen a los embeddings/Pinecone. Así, aunque por error se suba un PDF
    equivocado (ej. ficha de un alumno con DNI/telefono/cuenta), esa
    informacion nunca queda indexada ni puede salir en una respuesta.
    Devuelve el texto limpio + un resumen de cuantas coincidencias hubo por tipo.
    """
    clean_text = text
    summary = {}
    for label, pattern in SENSITIVE_PATTERNS.items():
        matches = pattern.findall(clean_text)
        if matches:
            summary[label] = len(matches)
            clean_text = pattern.sub(f"[DATO PROTEGIDO: {label}]", clean_text)
    return clean_text, summary


# ==========================================================================
# 2. EXTRACCION DE DOCX (texto + tablas, en orden)
# ==========================================================================

def extract_docx_blocks(path: str) -> List[str]:
    """Devuelve una lista de bloques: parrafos de texto y tablas en formato
    markdown, respetando el orden original del documento."""
    doc_file = DocxDocument(path)
    blocks = []
    for element in doc_file.element.body:
        if element.tag.endswith('}p'):
            para = Paragraph(element, doc_file)
            if para.text.strip():
                blocks.append(para.text)
        elif element.tag.endswith('}tbl'):
            table = Table(element, doc_file)
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                table_rows.append(f"| {' | '.join(cells)} |")
            blocks.append("[TABLA DETECTADA]\n" + "\n".join(table_rows))
    return blocks


# ==========================================================================
# 3. EXTRACCION DE PDF (texto + tablas + OCR de respaldo para paginas
#    escaneadas) usando pdfplumber en vez de PyPDFLoader.
#    PyPDFLoader solo saca texto plano y destruye la estructura de tablas;
#    pdfplumber permite extraer tablas como celdas reales.
# ==========================================================================

def _ocr_page(path: str, page_number: int) -> str:
    if not _preparar_ocr():
        return ""
    try:
        # poppler_path explícito: en Windows los binarios no siempre están en el
        # PATH del proceso que arranca el servidor, y sin esto convert_from_path
        # fallaba aunque Poppler estuviera instalado.
        extra = {"poppler_path": POPPLER_PATH} if POPPLER_PATH else {}
        images = _convert_from_path(
            path, first_page=page_number + 1, last_page=page_number + 1, dpi=200, **extra)
        if not images:
            return ""
        return _pytesseract.image_to_string(images[0], lang="spa+eng")
    except Exception as e:
        print(f"[OCR][WARN] No se pudo aplicar OCR a la pagina {page_number + 1}: {e}")
        return ""


def extract_pdf_blocks(path: str) -> List[str]:
    blocks = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = (page.extract_text() or "").strip()

            # Tablas: se extraen aparte para no perder su estructura de celdas
            tables = page.extract_tables()
            table_blocks = []
            for table in tables:
                if not table:
                    continue
                rows = [
                    f"| {' | '.join((cell or '').strip().replace(chr(10), ' ') for cell in row)} |"
                    for row in table
                ]
                table_blocks.append("[TABLA DETECTADA]\n" + "\n".join(rows))

            if not page_text and not table_blocks:
                # Pagina probablemente escaneada (imagen). Solo aqui gastamos OCR,
                # para minimizar costo/latencia en el resto del documento.
                ocr_text = _ocr_page(path, i).strip()
                if ocr_text:
                    blocks.append(ocr_text)
                continue

            if page_text:
                blocks.append(page_text)
            blocks.extend(table_blocks)
    return blocks


# ==========================================================================
# 4. CHUNKING CONSCIENTE DE TABLAS
#    Las tablas NUNCA se parten con el splitter generico de caracteres
#    (eso rompia filas a la mitad). Si una tabla es muy grande, se divide
#    por filas repitiendo la cabecera en cada fragmento para no perder
#    el significado de las columnas.
# ==========================================================================

TABLE_CHUNK_CHAR_LIMIT = 1500


def _split_large_table(table_block: str) -> List[str]:
    lines = table_block.split("\n")
    header = lines[0]  # "[TABLA DETECTADA]"
    rows = lines[1:]
    if not rows:
        return [table_block]

    header_row = rows[0]
    body_rows = rows[1:]

    chunks = []
    current = [header_row]
    current_len = len(header_row)
    for row in body_rows:
        if current_len + len(row) > TABLE_CHUNK_CHAR_LIMIT and len(current) > 1:
            chunks.append(header + " (parte)\n" + "\n".join(current))
            current = [header_row]
            current_len = len(header_row)
        current.append(row)
        current_len += len(row)
    if len(current) > 1:
        chunks.append(header + " (parte)\n" + "\n".join(current))
    return chunks if chunks else [table_block]


def build_documents_from_blocks(blocks: List[str], source: str, splitter: RecursiveCharacterTextSplitter) -> List[LangDocument]:
    final_docs: List[LangDocument] = []
    prose_buffer: List[str] = []

    def flush_prose():
        if not prose_buffer:
            return
        joined = "\n\n".join(prose_buffer)
        prose_buffer.clear()
        prose_doc = LangDocument(page_content=joined, metadata={"source": source})
        final_docs.extend(splitter.split_documents([prose_doc]))

    for block in blocks:
        if block.startswith("[TABLA DETECTADA]"):
            flush_prose()
            if len(block) > TABLE_CHUNK_CHAR_LIMIT:
                for part in _split_large_table(block):
                    final_docs.append(LangDocument(page_content=part, metadata={"source": source}))
            else:
                final_docs.append(LangDocument(page_content=block, metadata={"source": source}))
        else:
            prose_buffer.append(block)

    flush_prose()
    return final_docs


# ==========================================================================
# ENDPOINTS DE CONSULTA Y GESTION (sin cambios de logica)
# ==========================================================================

@router.get("/documents", response_model=List[ChatbotResponse])
def get_documents(db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    if current_user.get("rol") != "ADMIN":
        raise HTTPException(status_code=403, detail="No puedes ver esta información")
    return db.query(Chatbot).all()


@router.get("/download/{doc_id}")
def download_document(doc_id: int, db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    if current_user.get("rol") != "ADMIN":
        raise HTTPException(status_code=403, detail="No puedes ver esta información")
    doc = db.query(Chatbot).filter(Chatbot.id == doc_id).first()
    if not doc or not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="Archivo no encontrado.")
    return FileResponse(path=doc.file_path, filename=doc.filename)


@router.delete("/delete/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    if current_user.get("rol") != "ADMIN":
        raise HTTPException(status_code=403, detail="No puedes ver esta información")

    doc = db.query(Chatbot).filter(Chatbot.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado.")

    try:
        if doc.total_chunks and doc.total_chunks > 0:
            safe_filename = doc.filename.replace(" ", "_")
            ids_to_delete = [f"{safe_filename}-{i}" for i in range(doc.total_chunks)]
            pinecone_client.delete(ids_to_delete, namespace=NAMESPACE)
            print(f"[PINECONE] Se eliminaron {len(ids_to_delete)} chunks correctamente.")
    except Exception as e:
        print(f"[PINECONE WARNING] No se pudo borrar de Pinecone o ya no existe: {str(e)}")

    try:
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        db.delete(doc)
        db.commit()
        return {"message": f"'{doc.filename}' Eliminado correctamente del sistema."}
    except Exception as e:
        print(f"[ERROR LOCAL en /delete] {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error al eliminar los archivos locales o la base de datos: {str(e)}"
        )


# ==========================================================================
# ENDPOINT DE SUBIDA Y ENTRENAMIENTO
# ==========================================================================

@router.post("/upload")
async def upload(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: dict = Depends(get_current_user)):
    if current_user.get("rol") != "ADMIN":
        raise HTTPException(status_code=403, detail="No puedes ver esta información")

    count = db.query(Chatbot).count()
    if count >= MAX_DOCUMENTS:
        raise HTTPException(status_code=400, detail=f"Límite de {MAX_DOCUMENTS} archivos alcanzado.")

    existing = db.query(Chatbot).filter(Chatbot.filename == file.filename).first()
    if existing:
        raise HTTPException(
            status_code=400,
            detail=f"Ya existe un documento llamado '{file.filename}'. Elimínalo primero si quieres reemplazarlo."
        )

    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Formato no soportado. Usa PDF o DOCX.")

    content = await file.read()
    file_size = len(content)

    if file_size > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"El archivo excede el límite de {MAX_FILE_SIZE_MB}MB."
        )
    if file_size == 0:
        raise HTTPException(status_code=400, detail="El archivo está vacío.")

    unique_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    temp_path = os.path.join(UPLOAD_DIR, unique_name)
    try:
        with open(temp_path, "wb") as f:
            f.write(content)

        # 1. Extraccion por bloques (texto + tablas, con OCR de respaldo en PDF)
        if file.filename.lower().endswith((".docx", ".doc")):
            blocks = extract_docx_blocks(temp_path)
        elif file.filename.lower().endswith(".pdf"):
            blocks = extract_pdf_blocks(temp_path)
        else:
            os.remove(temp_path)
            return {"message": "Formato no soportado."}

        if not blocks:
            os.remove(temp_path)
            raise HTTPException(
                status_code=400,
                detail="El archivo no contiene texto legible (puede que sea una imagen sin OCR disponible o esté protegido)."
            )

        # 2. Redaccion de datos sensibles ANTES de indexar nada
        redacted_blocks = []
        redaction_summary = {}
        for block in blocks:
            clean_block, summary = redact_sensitive_data(block)
            redacted_blocks.append(clean_block)
            for k, v in summary.items():
                redaction_summary[k] = redaction_summary.get(k, 0) + v

        if redaction_summary:
            print(f"[SEGURIDAD] Se redactaron datos sensibles en '{file.filename}': {redaction_summary}")

        # 3. Chunking consciente de tablas
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,  # ~20% de solapamiento: evita cortar listas/tablas a la mitad
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        final_docs = build_documents_from_blocks(redacted_blocks, file.filename, text_splitter)

        if not final_docs:
            os.remove(temp_path)
            return {"message": "No se extrajo contenido del archivo."}

        safe_filename = file.filename.replace(" ", "_")
        ids = [f"{safe_filename}-{i}" for i in range(len(final_docs))]

        # 4. Subida por LOTES (reduce llamadas a la API de embeddings/Pinecone
        #    frente al 1-por-1 original). El upsert vía REST puede devolver
        #    200 OK sin que eso garantice al 100% que Pinecone ya lo persistió
        #    para lectura inmediata, así que tras cada lote se verifica
        #    CONTRA PINECONE DE VERDAD con /vectors/fetch y solo se reintentan
        #    los IDs que de verdad faltan.
        BATCH_SIZE = 10
        upserted_count = 0
        failed_ids: List[str] = []
        print(f"Iniciando subida de {len(final_docs)} chunks en lotes de {BATCH_SIZE}...")

        id_to_doc = dict(zip(ids, final_docs))

        def _vector_payload(doc_id: str, doc: LangDocument, vector: List[float]) -> dict:
            # "text" es la clave que usan los chunks ya indexados por
            # langchain_pinecone -- se mantiene el mismo nombre para que
            # /ask siga leyendo también los documentos indexados antes de
            # este cambio.
            return {
                "id": doc_id,
                "values": vector,
                "metadata": {"source": doc.metadata.get("source", file.filename), "text": doc.page_content},
            }

        for start in range(0, len(final_docs), BATCH_SIZE):
            batch_docs = final_docs[start:start + BATCH_SIZE]
            batch_ids = ids[start:start + BATCH_SIZE]

            try:
                batch_vectors = embeddings.embed_documents([d.page_content for d in batch_docs])
                payload = [_vector_payload(bid, doc, vec) for bid, doc, vec in zip(batch_ids, batch_docs, batch_vectors)]
                pinecone_client.upsert(payload, namespace=NAMESPACE)
            except Exception as batch_error:
                print(f"[PINECONE][WARN] Excepción al subir el lote {start}-{start+BATCH_SIZE}: {batch_error}")

            existing = pinecone_client.fetch_existing_ids(batch_ids, NAMESPACE)
            missing = [bid for bid in batch_ids if bid not in existing]

            if missing:
                print(f"[PINECONE][WARN] {len(missing)} chunk(s) no llegaron a Pinecone en el lote {start}-{start+BATCH_SIZE}. Reintentando 1 a 1...")
                for mid in list(missing):
                    try:
                        doc = id_to_doc[mid]
                        vec = embeddings.embed_documents([doc.page_content])[0]
                        pinecone_client.upsert([_vector_payload(mid, doc, vec)], namespace=NAMESPACE)
                    except Exception as chunk_error:
                        print(f"[ERROR Chunk {mid}] {str(chunk_error)}")

                # Verificación final de este lote tras los reintentos
                still_missing = [bid for bid in batch_ids if bid not in pinecone_client.fetch_existing_ids(batch_ids, NAMESPACE)]
                failed_ids.extend(still_missing)

        # Conteo final real (no estimado): cuántos IDs de este archivo existen
        # de verdad en Pinecone ahora mismo.
        confirmed_ids: set = set()
        for start in range(0, len(ids), 100):
            confirmed_ids |= pinecone_client.fetch_existing_ids(ids[start:start + 100], NAMESPACE)
        upserted_count = len(confirmed_ids)

        print(f"[PINECONE] Subida completada. Total confirmado en Pinecone: {upserted_count} de {len(final_docs)}")
        if failed_ids:
            print(f"[PINECONE][ALERTA] {len(failed_ids)} fragmento(s) no pudieron indexarse tras reintentos: {failed_ids[:20]}{'...' if len(failed_ids) > 20 else ''}")

        clean_ext = file_ext.replace(".", "")
        new_record = Chatbot(
            filename=file.filename,
            unique_filename=unique_name,
            file_path=temp_path,
            file_type=clean_ext,
            pinecone_index=INDEX_NAME,
            # IMPORTANTE: total_chunks debe ser len(final_docs) (el rango de IDs
            # que se generó, ej. "archivo-0" ... "archivo-N"), no el conteo
            # confirmado en Pinecone. /delete reconstruye los IDs a partir de
            # este número; si aquí guardáramos el conteo confirmado y hubo
            # huecos, /delete no podría borrar los IDs correctos.
            total_chunks=len(final_docs),
            status="entrenado"
        )
        db.add(new_record)
        db.commit()

        response = {
            "message": f"'{file.filename}' indexado correctamente.",
            "chunks_generados": len(final_docs),
            "chunks_confirmados_pinecone": upserted_count,
        }
        if failed_ids:
            response["advertencia_chunks_faltantes"] = (
                f"{len(failed_ids)} de {len(final_docs)} fragmentos no se pudieron confirmar en Pinecone "
                f"tras varios reintentos. Puede que falten partes de este documento en las respuestas del chatbot. "
                f"Puedes intentar eliminar y volver a subir el archivo."
            )
        if redaction_summary:
            response["advertencia_seguridad"] = (
                f"Se detectaron y protegieron datos sensibles en el documento: {redaction_summary}. "
                f"Si este archivo no debería contener ese tipo de información, revísalo y elimínalo."
            )
            response["datos_redactados"] = redaction_summary

        return response

    except HTTPException:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise
    except Exception as e:
        if os.path.exists(temp_path):
            os.remove(temp_path)
        print(f"ERROR REAL: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================================================
# ENDPOINT DE PREGUNTA (RAG)
# ==========================================================================

# --- SYSTEM PROMPT ---
# Se define UNA sola vez a nivel de módulo (no se reconstruye en cada request)
# y se envía como system_instruction en vez de concatenarlo al texto del
# usuario. Esto separa claramente "reglas" de "datos", y es más barato en
# tokens que meter todo en un único bloque de texto en cada llamada.
SYSTEM_PROMPT = """
Eres el asistente informativo oficial del colegio. Respondes de forma amigable
pero eres, ante todo, un motor de extracción de información analítico y
extremadamente riguroso. SOLO puedes usar la información que aparece en los
fragmentos de contexto (chunks) que se te entregan junto a cada pregunta.

REGLAS DE PROCESAMIENTO (ORDEN DE PRIORIDAD):

1. EXHAUSTIVIDAD OBLIGATORIA EN LISTAS Y TABLAS
   Si la pregunta requiere extraer requisitos, condiciones, faltas, normas,
   fechas o elementos de una lista, DEBES listar la TOTALIDAD de los elementos
   presentes en el contexto. Nunca te detengas en el primer ítem ni resumas u
   omitas puntos asumiendo que son secundarios. Si el contexto contiene 4
   requisitos o 2 faltas de máxima sanción, lista exactamente esos 4 y esas 2.

2. BÚSQUEDA INTEGRAL MULTI-CHUNK
   Revisa de forma transversal TODOS los fragmentos entregados antes de
   responder. Solo dices "No tengo esa información precisa 🏫" si, tras
   revisar hasta el último fragmento, confirmas que el dato no existe
   explícitamente. Si un dato está repartido en distintos fragmentos (ej. el
   horario de un nivel en un chunk y el de otro nivel en otro), únelos y
   sintetiza la respuesta.

3. DETECCIÓN DE ANOMALÍAS
   Si detectas inconsistencias numéricas, de fechas o normativas entre
   distintas secciones del contexto, señálalas explícitamente indicando qué
   dice cada sección, en vez de elegir una al azar.

4. AMBIGÜEDAD Y FECHAS (pagos/cuotas)
   Si el usuario pregunta por un pago pero hay varios (ej. una tabla con 12
   meses), pregunta primero a qué mes o cuota se refiere. Si pregunta "qué
   debo pagar", compara la fecha de hoy con las fechas de la tabla: si ya
   pasó, indica que está vencido; si falta poco, recuérdale el límite.

5. TABLAS
   La información puede venir en formato "| N° | PENSIÓN | FECHA |". Léela
   columna por columna, con cuidado.

6. MÚLTIPLES PREGUNTAS
   Si el usuario hace varias preguntas en un mismo mensaje, respóndelas todas,
   una por una, de forma clara y ordenada.

7. INFERENCIA CONTROLADA
   Si la respuesta no está textual pero puede deducirse combinando 2 o más
   fragmentos (ej. sumar cuotas, comparar fechas), hazlo y explica en una
   frase breve cómo llegaste a esa conclusión. No especules más allá de eso.

8. FIDELIDAD AL CONTEXTO
   Usa únicamente el contexto entregado. No inventes datos que no estén ahí.

9. PRIVACIDAD
   Nunca reveles números de DNI, teléfonos, correos personales, contraseñas o
   cuentas bancarias completas aunque aparezcan en el contexto. Si ves la
   marca "[DATO PROTEGIDO]" o algo que parezca un dato personal identificable,
   no lo repitas ni intentes reconstruirlo.

ESTILO Y FORMATO DE RESPUESTA:
- Usa un tono natural, directo, claro y conversacional.
- Presenta las respuestas estructuradas en párrafos breves o listas sencillas con viñetas (-).
- Utiliza negritas solo de forma sutil para destacar conceptos clave o nombres de apartados, evitando abusar de ellas en cada frase.
- NO utilices separadores horizontales (como "---" o "___").
- NO utilices encabezados de tipo Markdown pesado (como "###" o "##"). Si necesitas dividir secciones, hazlo con un título corto en negrita.
- No incluyas notas explicativas entre paréntesis ni aclaraciones meta sobre el documento al final, a menos que sean estrictamente necesarias.
- Ve directo al punto sin preámbulos innecesarios ("De acuerdo con la información...").
""".strip()

# --- CADENA DE MODELOS (ordenados por velocidad segun el benchmark) ---
# Si un modelo falla (error, 404, cuota agotada) o no arranca a tiempo,
# se pasa automaticamente al siguiente de la lista. Se dejan varios
# "escalones" (no solo un fallback) porque en el benchmark se vio que
# varios modelos individuales pueden fallar o quedarse sin cuota en
# cualquier momento (p. ej. gemini-flash-latest fallo en una corrida y en
# otra no) -- con 4 opciones en cadena es mucho mas dificil que el
# endpoint /ask se quede sin poder responder.
MODEL_CHAIN = [
    "gemini-flash-lite-latest",   # 1° mas rapido en el benchmark (~0.55s)
    "gemini-3.5-flash-lite",      # 2° mas rapido (~0.61s)
    "gemini-3.1-flash-lite",      # 3° mas rapido (~0.71s)
    "gemini-flash-latest",        # modelo "latest" estable, como ultimo respaldo
]

# Config de generación: temperatura 0 para respuestas factuales. Bajamos
# max_output_tokens de 1536 a 800: el tiempo de generación de un LLM escala
# ~linealmente con los tokens de SALIDA (no con el input), así que esto por
# sí solo reduce el tiempo de generación de forma proporcional. 800 tokens
# siguen siendo suficientes para listar exhaustivamente los ítems de un
# reglamento escolar (las reglas de "exhaustividad" del system prompt no
# cambian, solo se les da menos margen de relleno/verborrea).
GENERATION_CONFIG = genai_types.GenerateContentConfig(
    system_instruction=SYSTEM_PROMPT,
    temperature=0.0,
    max_output_tokens=2048,
)

# Con streaming, este timeout ya NO espera la respuesta completa de un
# modelo -- solo espera a que llegue el PRIMER fragmento de texto. Por eso
# se puede usar un valor agresivo: si en ~6s un modelo de la cadena ni
# siquiera empezó a responder, algo anda mal (sobrecarga, error, cuota) y
# tiene sentido pasar al siguiente. Si ya empezó a responder, seguimos
# streameando con ESE modelo hasta el final sin volver a chequear timeout.
# Al ULTIMO modelo de la cadena no se le aplica timeout (None): ya no hay
# a donde mas saltar, asi que es mejor esperar su respuesta completa que
# cortarlo a mitad de camino y dejar al usuario sin nada.
FIRST_CHUNK_TIMEOUT_SECONDS = 6
async def get_native_embedding_3072(text: str) -> List[float]:
    """
    Genera el vector de 3072 dimensiones usando el cliente ASÍNCRONO nativo (client.aio).
    Evita el overhead de hilos y reutiliza la conexión HTTP keep-alive.
    """
    response = await client.aio.models.embed_content(
        model="models/gemini-embedding-2",
        contents=text,
        config=genai_types.EmbedContentConfig(
            task_type="RETRIEVAL_QUERY",
            output_dimensionality=EMBEDDING_DIMENSIONS,
        )
    )
    # Accedemos a 'embeddings' (plural) y tomamos el primer resultado [0]
    return response.embeddings[0].values

def _stream_sync_to_queue(model: str, contents: str, q: "thread_queue.Queue") -> None:
    """Corre en un hilo aparte (nunca en el event loop). Lee el stream
    síncrono del SDK de Gemini y va empujando cada fragmento de texto a una
    cola thread-safe. Al terminar (bien o mal) empuja un centinela para que
    el lado async sepa que ya no hay más datos."""
    try:
        stream = client.models.generate_content_stream(
            model=model,
            contents=contents,
            config=GENERATION_CONFIG,
        )
        for chunk in stream:
            text = getattr(chunk, "text", None)
            if text:
                q.put(("chunk", text))
        q.put(("done", None))
    except Exception as e:
        q.put(("error", e))


async def _stream_model_response(
    model: str, contents: str, first_chunk_timeout: float = None
) -> AsyncGenerator[str, None]:
    """Generador async que va produciendo fragmentos de texto de un modelo
    de Gemini a medida que llegan (streaming real, no se espera la
    respuesta completa).

    Si `first_chunk_timeout` no es None y el PRIMER fragmento no llega
    dentro de ese tiempo, se levanta asyncio.TimeoutError -- esto permite
    que el llamador cambie de modelo (fallback) SIN haber emitido todavía
    nada al cliente, evitando respuestas mezcladas de dos modelos distintos.
    Una vez que el primer fragmento ya llegó, ya no se vuelve a aplicar
    timeout: se deja correr el streaming hasta el final.
    """
    q: "thread_queue.Queue" = thread_queue.Queue()
    thread = threading.Thread(target=_stream_sync_to_queue, args=(model, contents, q), daemon=True)
    thread.start()

    loop = asyncio.get_event_loop()
    first = True
    while True:
        get_call = loop.run_in_executor(None, q.get)
        if first and first_chunk_timeout is not None:
            kind, payload = await asyncio.wait_for(get_call, timeout=first_chunk_timeout)
        else:
            kind, payload = await get_call
        first = False

        if kind == "chunk":
            yield payload
        elif kind == "done":
            return
        elif kind == "error":
            raise payload


def _es_error_de_cuota(e: Exception) -> bool:
    """Detecta si la excepcion es por cuota agotada / rate limit (429,
    RESOURCE_EXHAUSTED) en vez de un fallo generico. Sirve solo para loggear
    con mas claridad -- la logica de fallback es la misma en ambos casos."""
    texto = str(e).upper()
    return "429" in texto or "RESOURCE_EXHAUSTED" in texto or "QUOTA" in texto


async def _ask_stream_generator(contents: str) -> AsyncGenerator[str, None]:
    """Orquesta la cadena MODEL_CHAIN para el endpoint /ask, streameando al
    cliente en cuanto hay texto disponible.

    - Si un modelo no arranca a tiempo, falla, devuelve 404, o se queda sin
      cuota ANTES de mandar el primer fragmento, se pasa al siguiente
      modelo de la cadena sin que el cliente note nada.
    - Si un modelo YA empezó a responder y falla a mitad de camino, NO se
      cambia de modelo (eso mezclaria dos respuestas distintas en un solo
      mensaje). En ese caso se corta el stream con un aviso breve.
    - El ultimo modelo de la cadena corre sin timeout de primer fragmento:
      ya no hay a donde mas saltar, asi que se le da todo el tiempo que
      necesite en vez de cortarlo antes de tiempo.
    """
    for idx, model in enumerate(MODEL_CHAIN):
        es_ultimo = idx == len(MODEL_CHAIN) - 1
        timeout = None if es_ultimo else FIRST_CHUNK_TIMEOUT_SECONDS
        t_model = time.perf_counter()
        ya_se_envio_algo = False
        try:
            async for piece in _stream_model_response(model, contents, first_chunk_timeout=timeout):
                ya_se_envio_algo = True
                yield piece
            print(f"[GEMINI] Respuesta completa de '{model}' en {time.perf_counter() - t_model:.2f}s.")
            return
        except Exception as e:
            elapsed = time.perf_counter() - t_model
            if isinstance(e, asyncio.TimeoutError):
                reason = "timeout esperando el primer fragmento"
            elif _es_error_de_cuota(e):
                reason = f"cuota agotada / rate limit ({type(e).__name__})"
            else:
                reason = f"{type(e).__name__}: {e}"

            if ya_se_envio_algo:
                # Este modelo ya habia mandado texto real al cliente: cambiar
                # de modelo ahora pegaria la respuesta de OTRO modelo detras
                # de un texto incompleto. Mejor cortar aca con un aviso claro
                # que devolver una respuesta incoherente.
                print(f"[GEMINI][ERROR] '{model}' falló A MITAD de la respuesta ({reason}) tras {elapsed:.2f}s. "
                      f"No se cambia de modelo (evita mezclar respuestas); se corta el stream.")
                yield "\n\n⚠️ Se interrumpió la respuesta por un problema técnico. Por favor, intenta de nuevo 🍎"
                return

            if es_ultimo:
                print(f"[GEMINI][ERROR REAL en /ask] El último modelo de la cadena ('{model}') también falló "
                      f"({reason}) tras {elapsed:.2f}s. Se agotó toda la cadena de modelos.")
                break

            print(f"[GEMINI][WARN] '{model}' falló ({reason}) tras {elapsed:.2f}s. "
                  f"Cambiando al siguiente modelo de la cadena (todavía no se envió nada al cliente).")
            continue

    yield "Error técnico o alta demanda del servicio, intenta de nuevo en unos segundos 🍎"


@router.post("/ask")
async def ask(question: str = Form(...)):
    t0 = time.perf_counter()
    hoy = datetime.now().strftime("%d de %B de %Y")

    # k ajustado (antes 4-6, ahora 5-7): el fallo típico en preguntas de
    # listas/normas ocurría porque el chunk relevante quedaba en la posición
    # 4-5 del ranking. Subir k mejora recall a un costo marginal en latencia.
    word_count = len(question.split())
    k = 5 if word_count <= 15 else 7

    # --------------------------------------------------------------------
    # Embedding y Pinecone medidos por separado (antes iban fusionados en
    # similarity_search). Similarity search en vez de MMR: para extracción
    # exhaustiva de listas y tablas normativas nos interesa recuperar los
    # chunks MÁS RELEVANTES, no los más "diversos" entre sí (que es justo lo
    # que optimiza MMR). MMR además es más lento porque calcula similitud
    # sobre fetch_k=30 candidatos antes de elegir k -- quitarlo mejora
    # precisión Y velocidad a la vez para este caso de uso.
    # --------------------------------------------------------------------
    t_embed = time.perf_counter()
    try:
        # Llamada directa async con timeout ajustado a 7s
        query_vector = await asyncio.wait_for(
            get_native_embedding_3072(question),
            timeout=20.0,
        )
    except asyncio.TimeoutError:
        print(
            f"[EMBEDDING][ERROR] Timeout de embedding (20s superado) tras {time.perf_counter() - t_embed:.2f}s"
        )
        return StreamingResponse(
            iter(
                [
                    "El servicio de embeddings tardó demasiado en responder, por favor reintenta 🍎"
                ]
            ),
            media_type="text/plain",
        )
    except Exception as e:
        print(
            f"[EMBEDDING][ERROR] Falló el embedding nativo tras {time.perf_counter() - t_embed:.2f}s: {type(e).__name__}: {e}"
        )
        return StreamingResponse(
            iter(
                [
                    "Error técnico al procesar la consulta, intenta de nuevo en unos segundos 🍎"
                ]
            ),
            media_type="text/plain",
        )

    embed_s = time.perf_counter() - t_embed

    t_pine = time.perf_counter()
    try:
        matches = await asyncio.to_thread(pinecone_client.query, query_vector, k, NAMESPACE)
    except Exception as e:
        print(f"[PINECONE][ERROR] Falló la búsqueda en el namespace '{NAMESPACE}' tras {time.perf_counter() - t_pine:.2f}s: {type(e).__name__}: {e}")
        return {"answer": "Error técnico o alta demanda del servicio, intenta de nuevo en unos segundos 🍎"}
    pinecone_s = time.perf_counter() - t_pine

    # "text"/"source" son las claves de metadata con las que se guardó cada
    # chunk en /upload (ver _vector_payload). Se reconstruye aquí una forma
    # mínima equivalente a los Document de langchain que se usaban antes.
    docs = [
        {"source": m.get("metadata", {}).get("source", "desconocido"), "text": m.get("metadata", {}).get("text", "")}
        for m in matches
    ]

    print(f"[PINECONE] Pregunta: '{question}' -> se recuperaron {len(docs)} chunks del namespace '{NAMESPACE}'.")
    if not docs:
        print("[PINECONE][ALERTA] No se recupero NINGUN chunk. Revisa que el namespace "
              "y la dimension del embedding coincidan con los documentos indexados.")

    contexto = "\n\n".join([f"FUENTE: {d['source']}\nCONTENIDO: {d['text']}" for d in docs])

    contents = f"""Fecha de hoy: {hoy}

CONTEXTO DEL REGLAMENTO:
{contexto}

PREGUNTA DEL USUARIO: {question}"""

    retrieval_elapsed = time.perf_counter() - t0
    print(f"[ASK] Retrieval listo en {retrieval_elapsed:.2f}s (embed={embed_s:.2f}s | pinecone={pinecone_s:.2f}s | "
          f"k={k}, chunks={len(docs)}). Iniciando streaming de la respuesta...")

    # StreamingResponse: el cliente empieza a recibir texto tan pronto como
    # el modelo primario emite su primer fragmento (normalmente 1-3s),
    # en vez de esperar la respuesta completa antes de recibir algo.
    # media_type text/plain funciona con fetch()/EventSource simples; si el
    # frontend espera SSE con formato "data: ...\n\n", envolver cada
    # `piece` en ese formato dentro de _ask_stream_generator.
    return StreamingResponse(_ask_stream_generator(contents), media_type="text/plain")