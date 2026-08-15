# -*- coding: utf-8 -*-
"""
generar_manual_auxiliar.py
Generador actualizado del Manual del Auxiliar en formato docx para I.E.P. Amancio Varona
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

DOC_PATH = r"c:\Users\jesus\OneDrive\Documentos\Proyecto Amancio Varona\Manuales\Manual_del_Auxiliar_Campus_Virtual.docx"

# Paleta de colores oficial
COLOR_GUINDA = RGBColor(0x70, 0x1C, 0x32)  # #701C32
COLOR_AZUL = RGBColor(0x09, 0x3E, 0x7A)    # #093E7A
COLOR_TEXTO = RGBColor(0x22, 0x22, 0x22)   # #222222
COLOR_GRIS_OSCURO = RGBColor(0x44, 0x44, 0x44)
COLOR_GRIS_MEDIO = RGBColor(0x66, 0x66, 0x66)
COLOR_GRIS_CLARO = RGBColor(0x88, 0x88, 0x88)

HEX_GUINDA = "701C32"
HEX_AZUL = "093E7A"
HEX_BG_GUINDA = "FDF2F4"
HEX_BG_AZUL = "F0F4F8"
HEX_BG_GRIS = "FAFAFA"
HEX_BORDER_DASHED = "9CA3AF"
HEX_BORDER_TABLE = "E5E7EB"


def crear_documento():
    doc = docx.Document()

    # Configuración de márgenes estándar (1 pulgada)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Estilos base
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXTO

    return doc


def add_p(doc, text="", style=None, align=None, space_after=6, bold=False, italic=False, color=None, size_pt=11):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(size_pt)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def add_bullet(doc, bold_prefix, text, space_after=4):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.font.name = 'Calibri'
        r_b.font.size = Pt(11)
        r_b.font.bold = True
        r_b.font.color.rgb = COLOR_TEXTO
    r_t = p.add_run(text)
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = COLOR_TEXTO
    return p


def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = COLOR_GUINDA
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = COLOR_AZUL
    return p


def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COLOR_GRIS_OSCURO
    return p


def add_callout(doc, title, text, theme="guinda"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)

    hex_bg = HEX_BG_GUINDA if theme == "guinda" else HEX_BG_AZUL
    hex_border = HEX_GUINDA if theme == "guinda" else HEX_AZUL

    # Shading and borders in XML
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_bg}" w:val="clear"/>')
    tcPr.append(shd)

    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{hex_border}"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

    mar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="160" w:type="dxa"/>
            <w:left w:w="220" w:type="dxa"/>
            <w:bottom w:w="160" w:type="dxa"/>
            <w:right w:w="200" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(mar)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r_title = p.add_run(title)
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_GUINDA if theme == "guinda" else COLOR_AZUL

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.15
    r_txt = p2.add_run(text)
    r_txt.font.name = 'Calibri'
    r_txt.font.size = Pt(10.5)
    r_txt.font.color.rgb = COLOR_TEXTO

    # Space after callout
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)


def add_captura_box(doc, num_captura, descripcion):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)

    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_BG_GRIS}" w:val="clear"/>')
    tcPr.append(shd)

    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="dashed" w:sz="6" w:space="0" w:color="{HEX_BORDER_DASHED}"/>
            <w:left w:val="dashed" w:sz="6" w:space="0" w:color="{HEX_BORDER_DASHED}"/>
            <w:bottom w:val="dashed" w:sz="6" w:space="0" w:color="{HEX_BORDER_DASHED}"/>
            <w:right w:val="dashed" w:sz="6" w:space="0" w:color="{HEX_BORDER_DASHED}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

    mar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="240" w:type="dxa"/>
            <w:left w:w="200" w:type="dxa"/>
            <w:bottom w:w="240" w:type="dxa"/>
            <w:right w:w="200" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(mar)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"CAPTURA {num_captura}")
    r1.font.name = 'Calibri'
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(f"[Pegar aquí: {descripcion}]")
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9.5)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)


def add_custom_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = col_widths[i]
        
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEX_GUINDA}" w:val="clear"/>')
        tcPr.append(shd)

        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.runs[0].font.name = 'Calibri'
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = col_widths[c_idx]

            tcPr = row_cells[c_idx]._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
            tcPr.append(shd)

            borders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_TABLE}"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_TABLE}"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_TABLE}"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER_TABLE}"/>
                </w:tcBorders>
            ''')
            tcPr.append(borders)

            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            if p.runs:
                p.runs[0].font.name = 'Calibri'
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = COLOR_TEXTO

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)


def construir_manual():
    doc = crear_documento()

    # PORTADA / ENCABEZADO
    add_p(doc, "I.E.P. AMANCIO VARONA — TUMÁN", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12, bold=True, color=COLOR_AZUL, space_after=2)
    add_p(doc, "Campus Virtual", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, color=COLOR_GRIS_MEDIO, space_after=6)
    add_p(doc, "MANUAL DEL AUXILIAR", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=24, bold=True, color=COLOR_GUINDA, space_after=4)
    add_p(doc, "Guía completa para el registro de asistencia diaria, gestión de partes disciplinarios y control de notas de conducta", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=13, color=COLOR_GRIS_OSCURO, space_after=4)
    add_p(doc, "Auxiliar de Educación · Año escolar 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11, color=COLOR_GRIS_CLARO, space_after=14)

    add_callout(
        doc,
        "Lo más importante de este manual",
        "Este manual abarca las funciones clave del Auxiliar: la asistencia diaria (Parte 3), el registro de partes disciplinarios (Parte 4) y la gestión de notas bimestrales de conducta (Parte 5). Ten en cuenta que los reportes de faltas no se pueden eliminar desde el campus para preservar la trazabilidad, mientras que las notas de conducta pueden supervisarse y ajustarse manualmente con confirmación en caso de discrepancias.",
        theme="guinda"
    )

    # PARTE 1. TU PRIMER INGRESO
    add_heading_1(doc, "Parte 1. Tu primer ingreso")

    add_heading_2(doc, "1.1 Cómo entrar al campus")
    add_bullet(doc, "1. ", "Abre tu navegador web de preferencia (Google Chrome, Microsoft Edge o Mozilla Firefox) e ingresa a la página web del colegio.")
    add_bullet(doc, "2. ", "Haz clic en el botón CAMPUS VIRTUAL ubicado en la parte superior derecha de la página.")
    add_bullet(doc, "3. ", "Escribe tu usuario y tu contraseña, y presiona el botón Iniciar Sesión.")

    add_captura_box(doc, 1, "Pantalla de acceso al campus virtual, con los campos de usuario y contraseña.")

    add_heading_3(doc, "Tu usuario y tu contraseña")
    add_p(doc, "Tu usuario es tu número de DNI. La primera vez que accedes al sistema, tu contraseña provisional también es tu número de DNI.")
    add_bullet(doc, "• Usuario: ", "Tu número de DNI (por ejemplo, 45781234).")
    add_bullet(doc, "• Contraseña inicial: ", "El mismo número de DNI.")

    add_heading_2(doc, "1.2 Lo que el campus te pedirá la primera vez")
    add_p(doc, "Al ingresar por primera vez, el campus te dirigirá automáticamente a una pantalla de seguridad obligatoria para proteger tu cuenta. Te solicitará dos datos fundamentales:")
    add_bullet(doc, "1. Contraseña nueva: ", "Debe contener al menos 8 caracteres (combina letras y números). No puede seguir siendo tu número de DNI.")
    add_bullet(doc, "2. Correo electrónico de contacto: ", "Para recibir comunicaciones institucionales y notificaciones del sistema.")

    add_captura_box(doc, 2, "Pantalla de primer ingreso solicitando cambio de contraseña y correo de contacto.")

    add_callout(
        doc,
        "Guarda bien tu contraseña",
        "Por medidas estrictas de seguridad, el campus no cuenta con recuperación pública por correo. Si olvidas tu contraseña, deberás solicitar a la Dirección o Administración del colegio el restablecimiento de tu acceso.",
        theme="azul"
    )

    add_heading_2(doc, "1.3 Cómo moverte por tu panel")
    add_p(doc, "En el lateral izquierdo de la pantalla dispones de una barra de navegación fija con los 5 módulos principales de tu función:")

    add_custom_table(
        doc,
        ["Sección", "Descripción y Uso"],
        [
            ["Inicio / Dashboard", "Resumen diario de actividades, accesos directos y recordatorios de tareas."],
            ["Asistencia", "Control y registro de asistencia diaria de estudiantes por nivel, grado y sección."],
            ["Reportes y Partes", "Registro de incidencias disciplinarias, faltas según reglamento y deméritos."],
            ["Notas de Conducta", "Supervisión bimestral de notas de conducta, cálculo automático y ajustes manuales."],
            ["Mensajería", "Canal de comunicación interna con los docentes del colegio."]
        ],
        [Inches(2.0), Inches(4.5)]
    )

    add_p(doc, "En dispositivos móviles o pantallas pequeñas, la barra lateral se oculta automáticamente y se abre tocando el icono de menú (tres líneas) en la esquina superior izquierda. Para cerrar sesión de forma segura, haz clic en el botón 'Cerrar Sesión' ubicado al pie de la barra.")

    # PARTE 2. LA PANTALLA DE INICIO
    add_heading_1(doc, "Parte 2. La pantalla de Inicio")
    add_p(doc, "Es la primera vista tras identificarte en el campus. Muestra tu saludo de bienvenida con tu nombre completo, la fecha oficial del día y tarjetas de acceso directo a los módulos de trabajo.")
    add_p(doc, "Asimismo, incluye avisos de tareas pendientes y la importancia de registrar la información de forma oportuna el mismo día, ya que la asistencia y los reportes alimentan en tiempo real el seguimiento psicopedagógico y administrativo.")

    add_captura_box(doc, 3, "Pantalla de Inicio del auxiliar con las tarjetas de acceso rápido y recordatorios.")

    # PARTE 3. CONTROL DE ASISTENCIA DIARIA
    add_heading_1(doc, "Parte 3. Control de asistencia diaria")
    add_p(doc, "Este módulo está diseñado para que el registro de asistencia de un aula completa tome menos de un minuto, optimizando la labor diaria del auxiliar.")

    add_heading_2(doc, "3.1 Los cuatro estados de asistencia")
    add_custom_table(
        doc,
        ["Letra", "Estado", "Criterio de Aplicación"],
        [
            ["P", "Presente", "El alumno asistió puntualmente. El sistema lo asigna por defecto a todos."],
            ["T", "Tardanza", "El alumno ingresó a la institución educativa después de la hora límite de entrada."],
            ["F", "Falta", "El alumno no asistió a clases y no cuenta con justificación presentada."],
            ["J", "Justificado", "El alumno no asistió, pero el apoderado presentó una justificación formal válida."]
        ],
        [Inches(0.8), Inches(1.5), Inches(4.2)]
    )

    add_callout(
        doc,
        "La diferencia entre Falta (F) y Justificado (J)",
        "El porcentaje de asistencia que visualiza el estudiante y su apoderado descuenta los días justificados del total de días lectivos. Marcar 'J' no perjudica su récord, mientras que 'F' reduce su porcentaje. Es vital registrar 'J' únicamente cuando exista el sustento correspondiente.",
        theme="azul"
    )

    add_heading_2(doc, "3.2 Cómo registrar la asistencia paso a paso")
    add_bullet(doc, "1. ", "Ingresa a 'Asistencia' desde la barra lateral izquierda.")
    add_bullet(doc, "2. ", "Verifica la fecha en la parte superior derecha. Por defecto aparece la fecha de hoy.")
    add_bullet(doc, "3. ", "Selecciona el Nivel (Primaria o Secundaria).")
    add_bullet(doc, "4. ", "Selecciona el Grado escolar correspondiente.")
    add_bullet(doc, "5. ", "Selecciona la Sección. La lista de estudiantes se cargará de manera automática y ordenada alfabéticamente.")
    add_bullet(doc, "6. ", "Todos los estudiantes inician marcados con 'P' (Presente). Haz clic únicamente sobre las letras T, F o J de aquellos alumnos que correspondan.")
    add_bullet(doc, "7. ", "Haz clic en el botón guinda 'Guardar Registro Diario' ubicado al final de la página.")

    add_captura_box(doc, 4, "Lista de asistencia de estudiantes con los selectores P, T, F, J y el contador en vivo.")

    add_heading_2(doc, "3.3 Contador en vivo y verificación")
    add_p(doc, "En la cabecera de la lista se actualizan en tiempo real cuatro contadores con puntos de color: Presentes, Tardanzas, Faltas y Justificados. Utilízalos como control visual previo al guardado (por ejemplo, si identificaste 2 faltas en el aula física, el contador debe marcar 2).")

    add_heading_2(doc, "3.4 Notificación automática a apoderados")
    add_p(doc, "Al presionar 'Guardar Registro Diario', el sistema procesa la información y despacha correos automáticos únicamente a los apoderados de los estudiantes que registraron Tardanza, Falta o Justificación, manteniendo informada a la familia de forma oportuna.")

    add_heading_2(doc, "3.5 Corrección de asistencia registrada")
    add_p(doc, "Si necesitas rectificar la asistencia de una fecha ya guardada, selecciona nuevamente la fecha y el aula, realiza las modificaciones requeridas y vuelve a presionar 'Guardar Registro Diario'. El sistema actualizará el registro existente sin duplicar datos.")

    # PARTE 4. REPORTES Y PARTES DE CONDUCTA
    add_heading_1(doc, "Parte 4. Reportes y partes de conducta")

    add_heading_2(doc, "4.1 Sistema de deméritos y semáforo disciplinario")
    add_p(doc, "Cada estudiante inicia el año escolar con 100 puntos de récord general. Cada reporte registrado descuenta puntos según la gravedad de la falta tipificada en el Reglamento Interno de la institución:")

    add_custom_table(
        doc,
        ["Puntaje Acumulado", "Semáforo", "Significado y Acción"],
        [
            ["100 a 75 puntos", "Verde · Buena conducta", "Estudiante sin observaciones disciplinarias relevantes."],
            ["74 a 40 puntos", "Amarillo · En observación", "Acumulación de faltas; requiere seguimiento preventivo."],
            ["39 a 0 puntos", "Rojo · Conducta crítica", "Caso crítico derivado a atención con Psicología y Dirección."]
        ],
        [Inches(1.8), Inches(2.2), Inches(2.5)]
    )

    add_p(doc, "Las faltas graves que conllevan 'Cambio de I.E.' colocan automáticamente al estudiante en estado crítico, independientemente de los puntos que conserve.")

    add_heading_2(doc, "4.2 Cómo registrar un nuevo reporte")
    add_bullet(doc, "1. ", "Entra a 'Reportes y Partes' y presiona el botón '+ Nuevo Reporte'.")
    add_bullet(doc, "2. ", "Busca al estudiante escribiendo al menos 3 letras de su nombre, apellido o DNI, y selecciónalo de la lista.")
    add_bullet(doc, "3. ", "Selecciona el Tipo de Falta (criterio general del reglamento).")
    add_bullet(doc, "4. ", "Selecciona la Falta específica. Se indicará entre paréntesis la cantidad de puntos a descontar.")
    add_bullet(doc, "5. ", "Verifica el recuadro descriptivo con el sustento reglamentario y la medida aplicable.")
    add_bullet(doc, "6. ", "Redacta la 'Descripción del suceso' detallando fecha, hora, lugar y hechos de forma objetiva (mínimo 10 caracteres).")
    add_bullet(doc, "7. ", "Presiona 'Registrar Reporte'.")

    add_captura_box(doc, 5, "Formulario modal de registro de nuevo reporte con selección de estudiante y falta reglamentaria.")

    add_heading_2(doc, "4.3 Resultado y notificaciones de alerta")
    add_p(doc, "Al registrar el reporte, el sistema muestra la tarjeta resumen con el nuevo puntaje del estudiante. Si cae en estado crítico o la falta involucra Cambio de I.E., se activará una alerta roja prioritaria para coordinar con Psicología y Dirección.")

    add_captura_box(doc, 6, "Tarjeta de confirmación con el puntaje resultante del alumno y las alertas pertinentes.")

    add_callout(
        doc,
        "Importancia de la redacción profesional",
        "La descripción registrada es visible en el expediente digital del estudiante y en el panel de Psicología. Debe redactarse con claridad, imparcialidad y apego a los hechos concretos.",
        theme="guinda"
    )

    add_heading_2(doc, "4.4 Bandeja e Historial de reportes")
    add_p(doc, "La pantalla principal muestra los últimos reportes registrados. Para consultar antecedentes completos de bimestres o años anteriores, utiliza el botón 'Historial de reportes' con buscador por estudiante.")

    add_captura_box(doc, 7, "Vista de historial de reportes con filtros y buscador de antecedentes disciplinarios.")

    # PARTE 5. GESTIÓN DE NOTAS DE CONDUCTA (NUEVA SECCIÓN)
    add_heading_1(doc, "Parte 5. Gestión de Notas de Conducta")
    add_p(doc, "Este módulo permite al auxiliar consultar, evaluar y asignar la calificación oficial bimestral de conducta de los estudiantes, garantizando total transparencia entre los reportes registrados y las notas que figurarán en las libretas de calificaciones.")

    add_heading_2(doc, "5.1 Cómo se calcula la nota bimestral de conducta")
    add_p(doc, "A diferencia del récord acumulado anual de 100 puntos, la calificación bimestral de conducta se evalúa en la escala vigesimal peruana (0 a 20):")
    add_bullet(doc, "• Base de partida: ", "Cada estudiante inicia cada bimestre con una nota base de 20 puntos.")
    add_bullet(doc, "• Reinicio bimestral: ", "Los puntos descontados por faltas se contabilizan dentro de las fechas límites de cada bimestre.")
    add_bullet(doc, "• Fórmula automática: ", "Nota Calculada = 20 − Total de puntos descontados por reportes en el bimestre (mínimo 0).")

    add_heading_2(doc, "5.2 Filtros de búsqueda y navegación")
    add_p(doc, "En la cabecera de la sección cuentas con un panel de filtros avanzados para segmentar la información:")
    add_bullet(doc, "• Año Escolar: ", "Selecciona el año académico lectivo (por defecto el año activo).")
    add_bullet(doc, "• Bimestre: ", "Elige entre el 1º, 2º, 3º o 4º Bimestre (inicia en el bimestre en curso).")
    add_bullet(doc, "• Nivel, Grado y Sección: ", "Filtra por aulas específicas para evaluar aula por aula.")
    add_bullet(doc, "• Buscador de Estudiante / DNI: ", "Escribe el nombre o documento para ubicar a un alumno de manera inmediata.")

    add_captura_box(doc, 8, "Pantalla de Notas de Conducta con filtros, métricas de resumen y la tabla de calificaciones.")

    add_heading_2(doc, "5.3 Estructura de la tabla de calificaciones")
    add_p(doc, "La tabla presenta las siguientes columnas informativas:")
    add_custom_table(
        doc,
        ["Columna", "Descripción y Utilidad"],
        [
            ["Estudiante", "Apellidos y nombres completos del estudiante junto a su número de DNI."],
            ["Sección", "Nivel educativo, grado y sección en la que se encuentra matriculado."],
            ["Incidencias", "Cantidad de reportes de conducta registrados durante el bimestre."],
            ["Puntos Restados", "Total de puntos descontados por faltas en el periodo (resaltado en rojo si > 0)."],
            ["Nota Calculada", "Calificación matemática según reglamento (Base 20 menos puntos descontados)."],
            ["Nota Final", "Casilla numérica editable donde figura la calificación efectiva del estudiante."],
            ["Origen / Estado", "Indica si la nota es 'Automática' o 'Manual', y si está 'Ajustada' respecto al cálculo."],
            ["Acciones", "Botón de restablecimiento para revertir cambios manuales al cálculo automático."]
        ],
        [Inches(1.8), Inches(4.7)]
    )

    add_heading_2(doc, "5.4 Modificación y guardado de nota manual")
    add_p(doc, "Para asignar o ajustar manualmente la nota de conducta de un estudiante:")
    add_bullet(doc, "1. ", "Haz clic sobre la casilla de 'Nota Final' del estudiante deseado.")
    add_bullet(doc, "2. ", "Ingresa la calificación deseada en el rango de 0 a 20 (admite decimales, ej. 16.5).")
    add_bullet(doc, "3. ", "Presiona la tecla Enter en tu teclado o haz clic en el botón de disco 'Guardar' que aparece al costado.")

    add_heading_2(doc, "5.5 Modal de confirmación por discrepancia")
    add_p(doc, "Si la nota ingresada manualmente no coincide con los puntos descontados por sus reportes de conducta (por ejemplo, el alumno tiene -4 puntos de falta que dan nota 16, pero ingresas 18 o viceversa), el sistema abrirá automáticamente una ventana modal de advertencia:")

    add_captura_box(doc, 9, "Ventana modal de confirmación por discrepancia al ingresar una nota que difiere de los reportes.")

    add_p(doc, "En este modal podrás revisar el resumen del estudiante, la cantidad de reportes, los puntos restados, la nota calculada por sistema y la nueva nota a asignar. Si confirmas que el cambio es intencional, presiona 'Sí, guardar nota manual' para registrarla con origen manual.")

    add_heading_2(doc, "5.6 Cómo restablecer la nota al cálculo automático")
    add_p(doc, "Si una nota fue modificada manualmente y deseas que vuelva a regirse estrictamente por los reportes registrados, haz clic en el botón circular 'Restablecer' (icono de flecha en retorno). El sistema borrará el ajuste manual y recalculará la nota en base a los deméritos del bimestre.")

    add_heading_2(doc, "5.7 Sincronización con Notas Finales y Libretas Oficiales")
    add_p(doc, "Toda nota de conducta visible en este módulo se sincroniza en tiempo real con la sábana de Notas Finales del Administrador y se imprime directamente en la libreta de calificaciones oficial del estudiante.")

    # PARTE 6. MENSAJERÍA
    add_heading_1(doc, "Parte 6. Mensajería")
    add_p(doc, "El campus cuenta con un canal de mensajería interna para coordinar actividades con los docentes de la institución:")
    add_bullet(doc, "• Inicio de conversaciones: ", "Puedes iniciar conversaciones directamente con cualquier docente.")
    add_bullet(doc, "• Recepción de mensajes: ", "Puedes responder mensajes enviados por administradores o directivos.")
    add_bullet(doc, "• Buscador de contactos: ", "Escribe el nombre del docente en el panel lateral de mensajería para abrir el chat.")

    add_captura_box(doc, 10, "Pantalla de mensajería interna con listado de docentes y panel de chat.")

    # PARTE 7. TU PERFIL, CONTRASEÑA Y NOTIFICACIONES
    add_heading_1(doc, "Parte 7. Tu perfil, tu contraseña y tus notificaciones")
    add_p(doc, "En la esquina superior derecha, al hacer clic sobre tu nombre, se despliega el menú de configuración personal:")
    add_bullet(doc, "• Mi perfil: ", "Permite revisar tus datos personales registrados y actualizar tu número telefónico o dirección.")
    add_bullet(doc, "• Cambiar contraseña: ", "Permite actualizar periódicamente tu clave secreta de acceso ingresando tu contraseña actual y la nueva de al menos 8 caracteres.")
    add_bullet(doc, "• Notificaciones: ", "Campana de avisos institucionales donde se alertan eventos o comunicados del colegio.")

    # PARTE 8. PREGUNTAS FRECUENTES
    add_heading_1(doc, "Parte 8. Preguntas frecuentes")

    add_heading_3(doc, "¿Olvidé mi contraseña, cómo puedo recuperarla?")
    add_p(doc, "Solicita a la Dirección o Secretaría del colegio el restablecimiento de tu contraseña. Por seguridad institucional, no se realiza recuperación automática por correo público.")

    add_heading_3(doc, "¿Puedo corregir una asistencia que ya guardé?")
    add_p(doc, "Sí. Selecciona la misma fecha, nivel, grado y sección, ajusta las marcas de los alumnos y vuelve a hacer clic en 'Guardar Registro Diario'. El sistema actualizará el registro.")

    add_heading_3(doc, "Un estudiante no tiene reportes de falta en el bimestre, ¿qué nota de conducta le corresponde?")
    add_p(doc, "Al no registrar deméritos, el sistema le asigna automáticamente la nota máxima de 20 (base 20 − 0 puntos = 20), visible tanto en tu panel como en las notas finales del administrador.")

    add_heading_3(doc, "Ingresé una nota de conducta distinta a la calculada y me salió un aviso en amarillo, ¿qué significa?")
    add_p(doc, "Es el control de discrepancia. Te avisa que la nota manual difiere de la calculada por sus faltas. Si estás seguro del cambio (por ejemplo, por evaluación integral o justificación de dirección), presiona 'Sí, guardar nota manual'.")

    add_heading_3(doc, "¿Registré un reporte con el estudiante equivocado, puedo borrarlo?")
    add_p(doc, "Los reportes no pueden borrarse directamente desde el campus para mantener la trazabilidad de faltas. Comunícalo a la Dirección del colegio para que se gestione la anulación del registro.")

    add_heading_3(doc, "¿Por qué el sistema no envía correos de asistencia a los alumnos Presentes?")
    add_p(doc, "Para evitar saturación en las bandejas de entrada de las familias y evitar que los correos sean clasificados como spam. Solo se notifica ante Tardanzas, Faltas y Justificaciones.")

    add_heading_3(doc, "¿Puedo usar el campus desde mi teléfono móvil o tablet?")
    add_p(doc, "Sí, el diseño es 100% responsivo y se adapta a cualquier pantalla. Para marcar listas extensas de asistencia o notas, se recomienda utilizar una computadora o tablet por mayor comodidad visual.")

    # SECCIÓN FINAL DE CONTACTO
    add_heading_1(doc, "¿Necesitas ayuda?")
    add_p(doc, "Si tienes dudas adicionales sobre el uso del campus virtual o requieres asistencia técnica, comunícate con el área de soporte tecnológico o con la Dirección de la I.E.P. Amancio Varona.")

    doc.save(DOC_PATH)
    print(f"Manual del Auxiliar guardado exitosamente en: {DOC_PATH}")


if __name__ == "__main__":
    construir_manual()
